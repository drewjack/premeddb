from flask import Flask
from flask_bootstrap import Bootstrap
from flask_sqlalchemy import SQLAlchemy
from flask import render_template, redirect, url_for, flash, send_file
from docx import Document
from io import BytesIO
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField, IntegerField, SelectField
from wtforms.validators import InputRequired, DataRequired, Email, Length
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_mail import Mail, Message
from flask import request
from itsdangerous import URLSafeTimedSerializer
import datetime
from flask_socketio import SocketIO, send, emit
from functools import wraps


# Flask app defined
app = Flask(__name__)

# Establishes secret key for app and email confirmation
app.config['SECRET_KEY'] = 'thisisasecret'
app.config['SECURITY_PASSWORD_SALT'] = 'alsoasecret'
ts = URLSafeTimedSerializer(app.config["SECRET_KEY"])
Bootstrap(app)

# Links to database which is created in config.py
app.config.from_pyfile('config.py')
db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Links to email account to send confirmation emails
# app.config['SECURITY_EMAIL_SENDER'] = 'no-reply@example.com'
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USERNAME'] = 'prehealthplanner@gmail.com'
app.config['MAIL_PASSWORD'] = 'supersonic'
mail = Mail(app)


# Establishes socketio
socketio = SocketIO(app)


class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(15), unique=True)
    firstname = db.Column(db.String(30))
    lastname = db.Column(db.String(30))
    email = db.Column(db.String(50), unique=True)
    year = db.Column(db.String(9), unique=False)
    appcycle = db.Column(db.String(20), unique=False)
    password = db.Column(db.String(80))
    admin = db.Column(db.String(1))
    urole = db.Column(db.String(80))
    confirmed = db.Column(db.BOOLEAN, nullable=False, default=False)
    scheduler = db.relationship('Scheduler', backref='student', lazy='dynamic')
    mcat = db.relationship('Mcat', backref='student', lazy='dynamic')
    grades = db.relationship('Grades', backref='student', lazy='dynamic')
    references = db.relationship('References', backref='student', lazy='dynamic')
    activities = db.relationship('Activities', backref='student', lazy='dynamic')
    status = db.relationship('Status', backref='student', lazy='dynamic')
    personal = db.relationship('Personal', backref='student', lazy='dynamic')
    post = db.relationship('Post', backref='student', lazy='dynamic')
    organizations = db.relationship('Organizations', backref='student', lazy='dynamic')

    def __repr__(self):
        return self.username

    def __init__(self, urole):
        self.urole = urole

    def get_urole(self):
            return self.urole


def login_required2(role="ANY"):
    def wrapper(fn):
        @wraps(fn)
        def decorated_view(*args, **kwargs):

            if not current_user.is_authenticated:
               return login_manager.unauthorized()
            urole = current_user.urole
            if ( (urole != role) and (role != "ANY")):
                return login_manager.unauthorized()
            return fn(*args, **kwargs)
        return decorated_view
    return wrapper


class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    recipient = db.Column(db.String(10000))
    title = db.Column(db.String(100))
    message = db.Column(db.String(10000))
    postdate = db.Column(db.String(100))


class Scheduler(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    schedulename = db.Column(db.String(50))
    schedule = db.Column(db.String(200))
    data = db.Column(db.LargeBinary)


class Grades(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    ogpa = db.Column(db.String(5))
    sgpa = db.Column(db.String(5))
    year = db.Column(db.String(20))

    def __repr__(self):
        return 'User %r' % (self.userid)


class Mcat(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    examdate = db.Column(db.String(15))
    overall = db.Column(db.String(5))
    cp = db.Column(db.String(5))
    cars = db.Column(db.String(5))
    bb = db.Column(db.String(5))
    ps = db.Column(db.String(5))

    def __repr__(self):
        return 'User %r' % (self.userid)


class References(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    name = db.Column(db.String(50))
    email = db.Column(db.String(50))
    type = db.Column(db.String(50))
    status = db.Column(db.String(500))


class Activities(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    activity = db.Column(db.String(50))
    type = db.Column(db.String(50))
    hours = db.Column(db.String(50))
    reference = db.Column(db.String(50))
    startdate = db.Column(db.String(50))
    enddate = db.Column(db.String(50))
    description = db.Column(db.String(10000))


class Status(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    university = db.Column(db.String(50))
    primary = db.Column(db.String(50))
    secondary = db.Column(db.String(50))
    interview = db.Column(db.String(50))
    offer = db.Column(db.String(50))
    essay1p = db.Column(db.String(500))
    essay1a = db.Column(db.String(10000))
    essay2p = db.Column(db.String(500))
    essay2a = db.Column(db.String(10000))
    essay3p = db.Column(db.String(500))
    essay3a = db.Column(db.String(10000))
    essay4p = db.Column(db.String(500))
    essay4a = db.Column(db.String(10000))
    essay5p = db.Column(db.String(500))
    essay5a = db.Column(db.String(10000))


class Personal(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    title = db.Column(db.String(50))
    essay = db.Column(db.String(10000))


class Medicalschools(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    medicalschool = db.Column(db.String(100))
    state = db.Column(db.String(100))
    city = db.Column(db.String(100))
    year = db.Column(db.String(100))
    wiki = db.Column(db.String(1000))
    homeurl = db.Column(db.String(1000))


class Organizations(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    userid = db.Column(db.Integer, db.ForeignKey('user.id'))
    name = db.Column(db.String(50))
    description = db.Column(db.String(400))
    type = db.Column(db.String(50))
    website = db.Column(db.String(400))
    facebook = db.Column(db.String(400))
    twitter = db.Column(db.String(400))
    email = db.Column(db.String(100))


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# Creates class for login form w/ username and password
class LoginForm(FlaskForm):
    # changed username to PeopleSoft number for now, would have to reset database and change required length
    username = StringField('PeopleSoft Number', validators=[InputRequired(), Length(min=4, max=15)])
    password = PasswordField('Password', validators=[InputRequired(), Length(min=8, max=80)])
    remember = BooleanField('Remember Me')


# Creates class for registration form w/ username, email, and password
class RegisterForm(FlaskForm):
    # changed username to PeopleSoft Number
    username = StringField('PeopleSoft Number', validators=[InputRequired(), Length(min=4, max=15)])
    firstname = StringField('First Name', validators=[InputRequired(), Length(min=4, max=15)])
    lastname = StringField('Last Name', validators=[InputRequired(), Length(min=4, max=15)])
    email = StringField('Email', validators=[InputRequired(), Email(message='Invalid email'), Length(max=50)])
    password = PasswordField('Password', validators=[InputRequired(), Length(min=8, max=80)])
    year = SelectField('Year', choices=[('Freshman', 'Freshman'), ('Sophomore', 'Sophomore'), ('Junior','Junior'), ('Senior','Senior')], validators=[InputRequired()])
    appcycle = SelectField('Which application cycle do you plan to apply?', choices=[('', ''), ('2018-2019','2018-2019'), ('2019-2020', '2019-2020'), ('2020-2021', '2020-2021'), ('2021-2022', '2021-2022'), ('2022-2023', '2022-2023'), ('2023-2024', '2023-2024'), ('2024-2025', '2024-2025'), ('2025-2026', '2025-2026'), ('2026-2027', '2026-2027')], validators=[InputRequired()])


# Creates classes for password reset
class EmailForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    peoplesoft = StringField('PeopleSoft Number', validators=[InputRequired(), Length(min=4, max=15)])


class PasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[DataRequired()])


# Creates classes for changes to account settings
class ChangeEmailForm(FlaskForm):
    email = StringField('New Contact Email', validators=[DataRequired(), Email()])


class ChangeYearForm(FlaskForm):
    year = SelectField('Year', choices=[('', ''), ('Freshman', 'Freshman'), ('Sophomore', 'Sophomore'), ('Junior', 'Junior'), ('Senior', 'Senior')], validators=[InputRequired()])


class ChangeAppCycleForm(FlaskForm):
    appcycle = SelectField('Which application cycle do you plan to apply?',
                           choices=[('', ''), ('2018-2019', '2018-2019'), ('2019-2020', '2019-2020'), ('2020-2021', '2020-2021'),
                                    ('2021-2022', '2021-2022'), ('2022-2023', '2022-2023'), ('2023-2024', '2023-2024'),
                                    ('2024-2025', '2024-2025'), ('2025-2026', '2025-2026'), ('2026-2027', '2026-2027')],
                           validators=[InputRequired()])


class ChangePasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[DataRequired()])


# Login Page
@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    global administrator
    global identifier
    form = LoginForm()
    failedlogin=""
    # Redirects to dashboard if already logged in
    if hasattr(current_user, 'id'):
        if current_user.id > 1:
            identifier = current_user.id
            return redirect(url_for('dashboard'))
    # Checks for user, password,and confirmation; otherwise gives error
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user:
            if check_password_hash(user.password, form.password.data):
                login_user(user, remember=form.remember.data)
                identifier = current_user.id
                if current_user.confirmed == True:
                    if current_user.admin == 'y':
                        administrator == 'y'
                        return redirect(url_for('administrator'))
                    else:
                        administrator == ''
                        return redirect(url_for('dashboard'))
                else:
                    failedlogin = "Account has not been activated yet. Please check your email."
            else:
                failedlogin = "Username or password is not valid"
        else:
            failedlogin = "Username or password is not valid"
    return render_template("login.html", form=form, failedlogin=failedlogin)


@app.route("/logout")
@login_required
def logout():
    logout_user()
    administrator == ''
    administratorstudent = 0
    return redirect(url_for('login'))


# Registration Page
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = RegisterForm()
    if form.validate_on_submit():
        hashed_password = generate_password_hash(form.password.data, method='sha256')
        new_user = User(username=form.username.data, firstname=form.firstname.data, lastname=form.lastname.data, email=form.email.data, year=form.year.data, appcycle=form.appcycle.data, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()
        token = generate_confirmation_token(new_user.email)
        confirm_url = url_for('confirm_email', token=token, _external=True)
        html = render_template('email.html', confirm_url=confirm_url)
        msg = Message('Registration Successful!', sender='prehealthplanner@gmail.com', recipients=[new_user.email])
        msg.body = html
        mail.send(msg)

        login_user(new_user)
        return redirect(url_for('newuser'))

    return render_template('signup.html', form=form)


def generate_confirmation_token(email):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(email, salt=app.config['SECURITY_PASSWORD_SALT'])


def confirm_token(token, expiration=3600):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        email = serializer.loads(
            token,
            salt=app.config['SECURITY_PASSWORD_SALT'],
            max_age=expiration
        )
    except:
        return False
    return email


# Post-registration page
@app.route('/newuser')
def newuser():
    return render_template("newuser.html")


# Post-confirmation page
@app.route('/confirm/<token>')
@login_required
def confirm_email(token):
    try:
        email = confirm_token(token)
    except:
        flash('The confirmation link is invalid or has expired.', 'danger')
    user = User.query.filter_by(email=email).first_or_404()
    if user.confirmed:
        flash('Account already confirmed. Please login.', 'success')
    else:
        user.confirmed = True
        db.session.add(user)
        db.session.commit()
        flash('You have confirmed your account. Thanks!', 'success')
    return redirect(url_for('login'))


# Password reset page
@app.route('/reset', methods=('GET', 'POST',))
def reset():
    form = EmailForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first_or_404()

        token = ts.dumps(user.email, salt='recover-key')

        recover_url = url_for(
            'reset_with_token',
            token=token,
            _external=True)

        html = render_template('static/additionalfiles/recover.html', recover_url=recover_url)
        msg = Message('Password Recovery', sender='prehealthplanner@gmail.com', recipients=[user.email])
        msg.body = html
        mail.send(msg)

        return render_template('forgotpassword.html')
    return render_template('reset.html', form=form)


@app.route('/reset/<token>', methods=["GET", "POST"])
def reset_with_token(token):
    try:
        email = ts.loads(token, salt="recover-key", max_age=86400)
    except:
        abort(404)

    form = PasswordForm()

    if form.validate_on_submit():
        user = User.query.filter_by(email=email).first_or_404()

        user.password = generate_password_hash(form.password.data, method='sha256')
        db.session.add(user)
        db.session.commit()

        return redirect(url_for('login'))

    return render_template('reset_with_token.html', form=form, token=token)


# Administrator functions

@app.route('/userdetails', methods=['POST', 'GET'])
@login_required
def userdetails():
    global administratorstudent
    global identifier
    administratorstudent = request.form['userdetails']
    # Allows admin to look into student account
    identifier = administratorstudent
    return redirect(url_for('dashboard'))


@app.route('/administrator', methods=['GET', 'POST'])
@login_required2(role="ADMIN")
def administrator():
    # Checks if you are an administrator or not
    # if current_user.admin != 'y':
    #     flash('You are not an administrator.')
    #     return redirect(url_for('dashboard'))
    return render_template("administrator.html")


@app.route('/studentrecords', methods=['GET', 'POST'])
@login_required
def studentrecords():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = User.query.all()
    if request.method == 'POST':
        result = User.query.filter(User.firstname.startswith(str(request.form['firstname'])), User.lastname.startswith(str(request.form['lastname'])), User.email.startswith(str(request.form['email'])))
    return render_template("studentrecords.html", result=result)


@app.route('/post', methods=['GET', 'POST'])
@login_required
def post():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id
    global freshman
    global sophomore
    global junior
    global senior
    global appcycle

    freshman = []
    sophomore = []
    junior = []
    senior = []
    appcycle = []
    recipientclass = []

    result = Post.query.all()
    if request.method == 'POST':
        title = request.form['title']
        message = request.form['message']
        now = datetime.datetime.now()
        postdate = now.strftime("%Y-%m-%d %I:%M %p")
        recipientlist = []
        if request.form.get('Freshman') == 'Freshman':
            freshman = User.query.filter_by(year='Freshman').all()
            recipientclass.append('Freshman')
        if request.form.get('Sophomore') == 'Sophomore':
            sophomore = User.query.filter_by(year="Sophomore").all()
            recipientclass.append('Sophomore')
        if request.form.get('Junior') == 'Junior':
            junior = User.query.filter_by(year="Junior").all()
            recipientclass.append('Junior')
        if request.form.get('Senior') == 'Senior':
            senior = User.query.filter_by(year="Senior").all()
            recipientclass.append('Senior')
        if request.form.get('appcycle') == 'appcycle':
            appcycle = User.query.filter_by(appcycle="2018-2019").all()
            recipientclass.append('Current App Cycle')

        recipientlist = freshman + sophomore + junior + senior + appcycle
        recipient = []
        for i in recipientlist:
            if i not in recipient:
                recipient.append(i.email)
        recipientstring = "['" + "','".join(recipient) + "']"
        fullname = current_user.firstname + ' ' + current_user.lastname
        signature = Post(userid=fullname, recipient=recipientstring, recipientclass=recipientclass, title=title, message=message, postdate=postdate)
        msg = Message('Pre-Health Portal Notification', sender='prehealthplanner@gmail.com', recipients=recipient)
        html = render_template("notificationemail.html", signature=signature)
        msg.body = html
        mail.send(msg)
        db.session.add(signature)
        db.session.commit()
        result = Post.query.all()
    return render_template("post.html", result=result)


@socketio.on('message')
def handleMessage(msg):
    print('Message: ' + msg)
    send(msg, broadcast=True)


@app.route('/editorganizations', methods=['POST', 'GET'])
@login_required
def editorganizations():
    global identifier
    # Failsafe to make sure identifier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = Organizations.query.all()
    if request.method == 'POST':
        name = request.form['name']
        description = request.form['description']
        type = request.form['type']
        website = request.form['website']
        facebook = request.form['facebook']
        twitter = request.form['twitter']
        email = request.form['email']
        signature = Organizations(name=name, description=description, type=type, website=website, facebook=facebook, twitter=twitter, email=email)
        db.session.add(signature)
        db.session.commit()
        result = Organizations.query.all()
    return render_template("editorganizations.html", result=result)


@app.route('/editorganizationsdetails', methods=['POST', 'GET'])
@login_required
def editorganizationsdetails():
    result = Organizations.query.filter_by(id=request.form['editorganizationsdetails']).first()
    return render_template("editorganizationsdetails.html", result=result)


@app.route('/editorganizationsprocess', methods=['POST', 'GET'])
def editorganizationsprocess():
    edit = Organizations.query.filter_by(id=request.form['update']).first()
    edit.description = request.form['description']
    edit.type = request.form['type']
    edit.website = request.form['website']
    edit.facebook = request.form['facebook']
    edit.twitter = request.form['twitter']
    edit.email = request.form['email']
    db.session.commit()
    return redirect(url_for('editorganizations'))


# Functions for all users
@app.context_processor
def notification():
    if current_user.is_authenticated:
        return dict(notification=Post.query.filter(Post.recipient.contains(current_user.email)).all())
    else:
        return dict(notification="")


@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    global identifier

    # Failsafe to make sure identifier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    if current_user.admin == 'y':
        admin = " (Admin: " + current_user.firstname + " " + current_user.lastname + " )"
    else:
        admin = ""

    user = User.query.filter_by(id=identifier).first()
    grades = Grades.query.filter_by(userid=identifier).all()
    result = Mcat.query.filter_by(userid=identifier).all()
    result1 = Activities.query.filter_by(userid=identifier).all()
    result2 = Post.query.all()
    result3 = []
    for dogs in result2:
        if user.year in dogs.recipient or user.appcycle in dogs.recipient:
            result3.append(dogs)
    ogpachart = [0,0,0,0,0,0,0,0]
    sgpachart = [0,0,0,0,0,0,0,0]
    cumchart = []
    pittgpa = [3.75,3.75,3.75,3.75,3.75,3.75,3.75,3.75]
    for item in grades:
        # Charts Graphs of Overall GPA
        if item.year == "1st Year Semester 1":
            ogpachart[0] = item.ogpa
        if item.year == "1st Year Semester 2":
            ogpachart[1] = item.ogpa
        if item.year == "2nd Year Semester 1":
            ogpachart[2] = item.ogpa
        if item.year == "2nd Year Semester 2":
            ogpachart[3] = item.ogpa
        if item.year == "3rd Year Semester 1":
            ogpachart[4] = item.ogpa
        if item.year == "3rd Year Semester 2":
            ogpachart[5] = item.ogpa
        if item.year == "4th Year Semester 1":
            ogpachart[6] = item.ogpa
        if item.year == "4th Year Semester 2":
            ogpachart[7] = item.ogpa
        # Charts Graphs of Science GPA
        if item.year == "1st Year Semester 1":
            sgpachart[0] = item.sgpa
        if item.year == "1st Year Semester 2":
            sgpachart[1] = item.sgpa
        if item.year == "2nd Year Semester 1":
            sgpachart[2] = item.sgpa
        if item.year == "2nd Year Semester 2":
            sgpachart[3] = item.sgpa
        if item.year == "3rd Year Semester 1":
            sgpachart[4] = item.sgpa
        if item.year == "3rd Year Semester 2":
            sgpachart[5] = item.sgpa
        if item.year == "4th Year Semester 1":
            sgpachart[6] = item.sgpa
        if item.year == "4th Year Semester 2":
            sgpachart[7] = item.sgpa
        # Graphs Cumulative GPA
        if item.year == "Cumulative":
            for i in range(8):
                cumchart.append(item.ogpa)
    #Ensures latest MCAT score is displayed
    iteration = 0
    pittmcat = [128.4,127.3,128.8,128.4]
    pittmcattotal = sum(map(int, pittmcat))
    mcat = ['0']
    mcattotal = 0
    for item1 in result:
        if item1.id > iteration:
            mcat = []
            mcat.append(item1.cp)
            mcat.append(item1.cars)
            mcat.append(item1.bb)
            mcat.append(item1.ps)
            mcattotal = sum(map(int, mcat))
            iteration = item1.id

    #Activities Tracking
    medvol = 0
    medemp = 0
    shadow = 0
    totvol = 0
    totres = 0
    totclin = 0
    for item2 in result1:
        if item2.type == "community service/volunteer – medical/clinical":
            medvol = medvol + int(item2.hours)
        if item2.type == "paid employment – medical/clinical":
            medemp = medemp + int(item2.hours)
        if item2.type == "physician shadowing/clinical observation":
            shadow = shadow + int(item2.hours)
        if item2.type == "community service/volunteer – medical/clinical" or item2.type == "community service/volunteer – not medical/clinical":
            totvol = totvol + int(item2.hours)
        if item2.type == "research/lab" or item2.type == "presentations/posters":
            totres = totres + int(item2.hours)
        if item2.type == "physician shadowing/clinical observation" or item2.type == "community service/volunteer – medical/clinical" or item2.type == "paid employment – medical/clinical":
            totclin = totclin + int(item2.hours)
    return render_template("dashboard.html", user=user, admin=admin, name=current_user.username, pittgpa=pittgpa, ogpachart=ogpachart, sgpachart=sgpachart, cumchart=cumchart, pittmcat=pittmcat, pittmcattotal=pittmcattotal, mcat=mcat, mcattotal=mcattotal, medvol=medvol, medemp=medemp, shadow=shadow, totvol=totvol, totres=totres, totclin=totclin, result3=result3)


@app.route('/scheduler', methods=['GET', 'POST'])
@login_required
def scheduler():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = Scheduler.query.filter_by(userid=identifier).all()
    if request.method == 'POST':
        if int(Scheduler.query.count()) < 6:
            file = request.files['inputfile']
            schedulename = request.form['schedulename']
            data = file.read()
            signature = Scheduler(userid=identifier, schedulename=schedulename,schedule=file.filename, data=data)
            db.session.add(signature)
            db.session.commit()
            result = Scheduler.query.filter_by(userid=identifier).all()
    return render_template("scheduler.html", result=result)


@login_required
@app.route('/schedulerdownload', methods=['GET', 'POST'])
def schedulerdownload():
    file_data = Scheduler.query.filter_by(id=request.form['schedulerdownload']).first()
    return send_file(BytesIO(file_data.data), attachment_filename=file_data.schedule, as_attachment=True)


@app.route('/academics', methods=['POST', 'GET'])
@login_required
def academics():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    grades = Grades.query.filter_by(userid=identifier).all()
    result = Mcat.query.filter_by(userid=identifier).all()
    result1 = References.query.filter_by(userid=identifier).all()
    if request.method == 'POST':
        ogpa = request.form['ogpa']
        sgpa = request.form['sgpa']
        year = request.form['year']
        signature = Grades(userid=identifier, ogpa=ogpa, sgpa=sgpa, year=year)
        db.session.add(signature)
        db.session.commit()
        grades = Grades.query.filter_by(userid=identifier).all()
    return render_template("academics.html", result=result, result1=result1, grades=grades)


@app.route('/academicsdetails', methods=['POST', 'GET'])
@login_required
def academicsdetails():
    result = References.query.filter_by(id=request.form['academicsdetails']).first()
    return render_template("academicsdetails.html", result=result)


@app.route('/academicsdetailsprocess', methods=['POST', 'GET'])
@login_required
def academicsdetailsprocess():
    edit = References.query.filter_by(id=request.form['update']).first()
    edit.email = request.form['email']
    edit.type = request.form['type']
    edit.status = request.form['status']
    db.session.commit()
    return redirect(url_for('academics'))


@app.route('/mcat', methods=['POST', 'GET'])
@login_required
def mcat():
    if request.method == 'POST':
        if request.form['examdate'] != '':
            examdate = request.form['examdate']
            overall = request.form['overall']
            cp = request.form['cp']
            cars = request.form['cars']
            bb = request.form['bb']
            ps = request.form['ps']
            signature = Mcat(userid=identifier, examdate=examdate, overall=overall, cp=cp, cars=cars, bb=bb, ps=ps)
            db.session.add(signature)
            db.session.commit()
    return redirect(url_for('academics'))


@app.route('/references', methods=['POST', 'GET'])
@login_required
def references():
    if request.method == 'POST':
        if request.form['name'] != '':
            name = request.form['name']
            email = request.form['email']
            type = request.form['type']
            status = request.form['status']
            signature = References(userid=identifier, name=name, email=email, type=type, status=status)
            db.session.add(signature)
            db.session.commit()
    return redirect(url_for('academics'))


@app.route('/activities', methods=['POST', 'GET'])
@login_required
def activities():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = Activities.query.filter_by(userid=identifier).all()
    if request.method == 'POST':
        if request.form['activity'] != '':
            activity = request.form['activity']
            type = request.form['type']
            reference = request.form['reference']
            hours = request.form['hours']
            signature = Activities(userid=identifier, activity=activity, type=type, reference=reference, hours=hours)
            db.session.add(signature)
            db.session.commit()
            result = Activities.query.filter_by(userid=identifier).all()
    return render_template("activities.html", result=result)


@app.route('/activitiesdetails', methods=['POST', 'GET'])
@login_required
def activitiesdetails():
    result = Activities.query.filter_by(id=request.form['activitiesdetails']).first()
    return render_template("activitiesdetails.html", result=result)


@app.route('/activitiesdetailsprocess', methods=['POST', 'GET'])
@login_required
def activitiesdetailsprocess():
    edit = Activities.query.filter_by(id=request.form['update']).first()
    edit.type = request.form['type']
    edit.hours = request.form['hours']
    edit.reference = request.form['reference']
    edit.startdate = request.form['startdate']
    edit.enddate = request.form['enddate']
    edit.description = request.form['description']
    db.session.commit()
    return redirect(url_for('activities'))


@app.route('/status', methods=['POST', 'GET'])
@login_required
def status():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = Status.query.filter_by(userid=identifier).all()
    medicalschools = Medicalschools.query.all()
    if request.method == 'POST':
        if request.form['university'] != '':
            university = request.form['university']
            primary = request.form['primary']
            secondary = request.form['secondary']
            interview = request.form['interview']
            offer = request.form['offer']
            signature = Status(userid=identifier, university=university, primary=primary, secondary=secondary, interview=interview, offer=offer)
            db.session.add(signature)
            db.session.commit()
    return render_template("status.html", result=result, medicalschools=medicalschools)


@app.route('/statusdetails', methods=['POST', 'GET'])
@login_required
def statusdetails():
    result = Status.query.filter_by(id=request.form['statusdetails']).first()
    medicalschool = Medicalschools.query.filter_by(medicalschool=result.university).first()
    return render_template("statusdetails.html", result=result, medicalschool=medicalschool)


@app.route('/statusdetailsprocess', methods=['POST', 'GET'])
@login_required
def statusdetailsprocess():
    edit = Status.query.filter_by(id=request.form['update']).first()
    edit.primary = request.form['primary']
    edit.secondary = request.form['secondary']
    edit.interview = request.form['interview']
    edit.offer = request.form['offer']
    edit.essay1p = request.form['essay1p']
    edit.essay1a = request.form['essay1a']
    edit.essay2p = request.form['essay2p']
    edit.essay2a = request.form['essay2a']
    edit.essay3p = request.form['essay3p']
    edit.essay3a = request.form['essay3a']
    edit.essay4p = request.form['essay4p']
    edit.essay4a = request.form['essay4a']
    edit.essay5p = request.form['essay5p']
    edit.essay5a = request.form['essay5a']
    db.session.commit()
    return redirect(url_for('status'))


@app.route('/statusdetailsword', methods=['POST', 'GET'])
@login_required
def statusdetailsword():
    edit = Status.query.filter_by(id=request.form['word']).first()
    document = Document()
    document.add_heading(edit.university, 0)
    #Essay 1
    document.add_heading('Prompt 1:', level=2)
    document.add_paragraph(edit.essay1p)
    document.add_heading('Essay 1:', level=2)
    document.add_paragraph(edit.essay1a)
    #Essay 2
    document.add_heading('Prompt 2:', level=2)
    document.add_paragraph(edit.essay2p)
    document.add_heading('Essay 2:', level=2)
    document.add_paragraph(edit.essay2a)
    #Essay 3
    document.add_heading('Prompt 3:', level=2)
    document.add_paragraph(edit.essay3p)
    document.add_heading('Essay 3:', level=2)
    document.add_paragraph(edit.essay3a)
    #Essay 4
    document.add_heading('Prompt 4:', level=2)
    document.add_paragraph(edit.essay4p)
    document.add_heading('Essay 4:', level=2)
    document.add_paragraph(edit.essay4a)
    #Essay 5
    document.add_heading('Prompt 5:', level=2)
    document.add_paragraph(edit.essay5p)
    document.add_heading('Essay 5:', level=2)
    document.add_paragraph(edit.essay5a)
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename='report.doc')


@app.route('/personalstatement', methods=['POST', 'GET'])
@login_required
def personalstatement():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id

    result = Personal.query.filter_by(userid=identifier).all()
    if request.method == 'POST':
        if request.form['title'] != '':
            title = request.form['title']
            signature = Personal(userid=identifier, title=title)
            db.session.add(signature)
            db.session.commit()
            result = Personal.query.filter_by(userid=identifier).all()
    return render_template("personalstatement.html", result=result)


@app.route('/personalstatementdetails', methods=['POST', 'GET'])
def personalstatementdetails():
    result = Personal.query.filter_by(id=request.form['personalstatementdetails']).first()
    return render_template("personalstatementdetails.html", result=result)


@app.route('/personalstatementdetailsprocess', methods=['POST', 'GET'])
@login_required
def personalstatementdetailsprocess():
    edit = Personal.query.filter_by(id=request.form['update']).first()
    edit.essay = request.form['essay']
    db.session.commit()
    return redirect(url_for('personalstatement'))


@app.route('/personaldetailsword', methods=['POST', 'GET'])
@login_required
def personaldetailsword():
    studentname = User.query.filter_by(id=identifier).first()
    edit = Personal.query.filter_by(id=request.form['word']).first()
    filename = str(studentname.firstname) + str(studentname.lastname) + str(edit.title) + 'personalstatement.doc'
    document = Document()
    document.add_heading('Personal Statement: ' + edit.title, 0)
    document.add_paragraph(edit.essay)
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename=filename)


@app.route('/information', methods=['POST', 'GET'])
@login_required
def information():
    return render_template("information.html")

@app.route('/advisors', methods=['POST', 'GET'])
@login_required
def advisors():
    return render_template("advisors.html")

@app.route('/organizations', methods=['POST', 'GET'])
@login_required
def organizations():
    global identifier
    # Failsafe to make sure idenfitier is set to current user if not previously set
    try:
        identifier
    except:
        identifier = current_user.id
    result = Organizations.query.all()
    return render_template("organizations.html", result=result)

@app.route('/universitystats', methods=['POST', 'GET'])
@login_required
def universitystats():
    global identifier
    result = Post.query.all()
    if request.method == 'POST':
        title = request.form['title']
        message = request.form['message']
        now = datetime.datetime.now()
        postdate = now.strftime("%Y-%m-%d %I:%M %p")
        recipient = []
        if request.form.get('Freshman') == 'Freshman':
            recipient.append('Freshman')
        if request.form.get('Sophomore') == 'Sophomore':
            recipient.append('Sophomore')
        if request.form.get('Junior') == 'Junior':
            recipient.append('Junior')
        if request.form.get('Senior') == 'Senior':
            recipient.append('Senior')
        if request.form.get('appcycle') == 'appcycle':
            recipient.append('appcycle')
        recipientstring = "['" + "','".join(recipient) + "']"
        fullname = current_user.firstname + ' ' + current_user.lastname
        signature = Post(userid=fullname, recipient=recipientstring, title=title, message=message, postdate=postdate)
        db.session.add(signature)
        db.session.commit()
        result = Post.query.all()
    return render_template("universitystats.html", result=result)


@app.route('/settings', methods=['POST', 'GET'])
@login_required
def settings():
    user = User.query.filter_by(id=current_user.id).first()
    return render_template("settings.html", user=user, email=current_user.email, year=current_user.year, appcycle=current_user.appcycle)


@app.route('/changeemail', methods=['POST', 'GET'])
@login_required
def changeemail():
    form = ChangeEmailForm()
    if form.validate_on_submit():
        user = User.query.filter_by(id=current_user.id).first()
        user.email = form.email.data
        db.session.add(user)
        db.session.commit()
    return render_template("changeemail.html", form=form)


@app.route('/changeyear', methods=['POST', 'GET'])
@login_required
def changeyear():
    form = ChangeYearForm()
    if form.validate_on_submit():
        user = User.query.filter_by(id=current_user.id).first()
        user.year = form.year.data
        db.session.add(user)
        db.session.commit()
    return render_template("changeyear.html", form=form)


@app.route('/changeappcycle', methods=['POST', 'GET'])
@login_required
def changeappcycle():
    form = ChangeAppCycleForm()
    if form.validate_on_submit():
        user = User.query.filter_by(id=current_user.id).first()
        user.year = form.appcycle.data
        db.session.add(user)
        db.session.commit()
    return render_template("changeappcycle.html", form=form)


@app.route('/changepassword', methods=["GET", "POST"])
@login_required
def changepassword():
    failedlogin = ""
    form = ChangePasswordForm()

    if form.validate_on_submit():
        user = User.query.filter_by(id=current_user.id).first()
        user.password = generate_password_hash(form.password.data, method='sha256')
        db.session.add(user)
        db.session.commit()
    return render_template("changepassword.html", form=form, failedlogin=failedlogin)


# Deletion Routes


@app.route('/deleteaccount', methods=['POST', 'GET'])
@login_required
def deleteaccount():
    User.query.filter_by(id=current_user.id).delete()
    Scheduler.query.filter_by(userid=current_user.id).delete()
    Grades.query.filter_by(userid=current_user.id).delete()
    Mcat.query.filter_by(userid=current_user.id).delete()
    References.query.filter_by(userid=current_user.id).delete()
    Activities.query.filter_by(userid=current_user.id).delete()
    Status.query.filter_by(userid=current_user.id).delete()
    Personal.query.filter_by(userid=current_user.id).delete()
    db.session.commit()
    return redirect(url_for('login'))


@app.route('/deleteuser', methods=['POST', 'GET'])
@login_required
def deleteuser():
    User.query.filter_by(id=int(request.form['userdelete'])).delete()
    Scheduler.query.filter_by(userid=int(request.form['userdelete'])).delete()
    Grades.query.filter_by(userid=int(request.form['userdelete'])).delete()
    Mcat.query.filter_by(userid=int(request.form['userdelete'])).delete()
    References.query.filter_by(userid=int(request.form['userdelete'])).delete()
    Activities.query.filter_by(userid=int(request.form['userdelete'])).delete()
    Status.query.filter_by(userid=int(request.form['userdelete'])).delete()
    Personal.query.filter_by(userid=int(request.form['userdelete'])).delete()
    db.session.commit()
    return redirect(url_for('studentrecords'))


@app.route('/deletepost', methods=['POST', 'GET'])
@login_required
def deletepost():
    if request.form['postdelete'] != '':
        Post.query.filter_by(id=int(request.form['postdelete'])).delete()
        db.session.commit()
    return redirect(url_for('post'))


@app.route('/deleteorganizations', methods=['POST', 'GET'])
@login_required
def deleteorganizations():
    if request.form['organizationsdelete'] != '':
        Organizations.query.filter_by(id=int(request.form['organizationsdelete'])).delete()
        db.session.commit()
    return redirect(url_for('editorganizations'))


@app.route('/deletescheduler', methods=['POST', 'GET'])
@login_required
def deletescheduler():
    if request.form['schedulerdelete'] != '':
        Scheduler.query.filter_by(id=int(request.form['schedulerdelete'])).delete()
        db.session.commit()
    return redirect(url_for('scheduler'))


@app.route('/deletegrades', methods=['POST', 'GET'])
@login_required
def deletegrades():
    if request.form['gradesdelete'] != '':
        Grades.query.filter_by(id=int(request.form['gradesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deletemcat', methods=['POST', 'GET'])
@login_required
def deletemcat():
    if request.form['mcatdelete'] != '':
        Mcat.query.filter_by(id=int(request.form['mcatdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deletereferences', methods=['POST', 'GET'])
@login_required
def deletereferences():
    if request.form['referencesdelete'] != '':
        References.query.filter_by(id=int(request.form['referencesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('academics'))


@app.route('/deleteactivities', methods=['POST', 'GET'])
@login_required
def deleteactivities():
    if request.form['activitiesdelete'] != '':
        Activities.query.filter_by(id=int(request.form['activitiesdelete'])).delete()
        db.session.commit()
    return redirect(url_for('activities'))


@app.route('/deletestatus', methods=['POST', 'GET'])
def deletestatus():
    if request.form['statusdelete'] != '':
        Status.query.filter_by(id=int(request.form['statusdelete'])).delete()
        db.session.commit()
    return redirect(url_for('status'))


@app.route('/deletepersonalstatement', methods=['POST', 'GET'])
@login_required
def deletepersonalstatement():
    if request.form['personalstatementdelete'] != '':
        Personal.query.filter_by(id=int(request.form['personalstatementdelete'])).delete()
        db.session.commit()
    return redirect(url_for('personalstatement'))


# Makes Summary Word Doc
@app.route('/summary')
@login_required
def summary():
    activities = Activities.query.filter_by(userid=identifier).all()
    grades = Grades.query.filter_by(userid=identifier).all()
    mcat = Mcat.query.filter_by(userid=identifier).all()
    references = References.query.filter_by(userid=identifier).all()
    status = Status.query.filter_by(userid=identifier).all()
    studentname = User.query.filter_by(id=identifier).first()
    filename = str(studentname.firstname) + str(studentname.lastname) + 'summary.doc'

    document = Document()
    document.add_heading("Summary", 0)
    document.add_heading('GPA', 1)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Overall'
    hdr_cells[1].text = 'Science'
    for item in grades:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.ogpa)
        row_cells[1].text = str(item.sgpa)

    document.add_heading('MCAT', 1)
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Overall'
    hdr_cells[2].text = 'C/P'
    hdr_cells[3].text = 'CARS'
    hdr_cells[4].text = 'B/B'
    hdr_cells[5].text = 'P/S'
    for item in mcat:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.examdate)
        row_cells[1].text = str(item.overall)
        row_cells[2].text = str(item.cp)
        row_cells[3].text = str(item.cars)
        row_cells[4].text = str(item.bb)
        row_cells[5].text = str(item.ps)

    document.add_heading('References', 1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Status'
    for item in references:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
        row_cells[1].text = str(item.email)
        row_cells[2].text = str(item.type)
        row_cells[3].text = str(item.status)

    document.add_heading('Activities', 1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Activity'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Reference'
    hdr_cells[3].text = 'Hours'
    for item in activities:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.activity)
        row_cells[1].text = str(item.type)
        row_cells[2].text = str(item.reference)
        row_cells[3].text = str(item.hours)

    document.add_heading('Application Status', 1)
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'University'
    hdr_cells[1].text = 'Primary'
    hdr_cells[2].text = 'Secondary'
    hdr_cells[3].text = 'Interview'
    hdr_cells[4].text = 'Offer'
    for item in status:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.university)
        row_cells[1].text = str(item.primary)
        row_cells[2].text = str(item.secondary)
        row_cells[3].text = str(item.interview)
        row_cells[4].text = str(item.offer)

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    return send_file(f, as_attachment=True, attachment_filename=filename)


# Flask app initialized
if __name__ == '__main__':
    app.run()
